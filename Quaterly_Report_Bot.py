from docx import Document
import os

def get_docx_files(path):
    return [f for f in os.listdir(path) if f.endswith('.docx')]

def load_docx(path, filename):
    return Document(path + filename)

def create_new_master_docx(year, qtr):
    return Document(), f'{year}_{qtr}_Master.docx'

def save_master_docx(path, filename, document):
    document.save(path + filename)

def append_pragraph_to_docx(master, person, document, section):
    # Locate the section
    section = (section * 3)+ 1

    # Locate the section answer
    text = document.tables[0].cell(section, 0).text

    # remove 'Type here' text from the document
    text = text.replace('Type here\n\n', '')
    text = text.replace('Type here\n', '')

    if text in ['']:
        return
    
    # person in bold
    p = master.add_paragraph()
    p.add_run(person).bold = True
    # add the first cell in the table from document to master
    master.add_paragraph(text)
    # add a blank line
    master.add_paragraph()





if __name__ == '__main__':
    year = 2023
    qtr = 'Q2'
    fileLocation = 'C:/Users/Jeff/Downloads/2023Q2/'
    sections = ["Publications:", "Conference Presentations:", "New Tools, Methods & Datasets:", "Impact and Uptake:",
        "Capability and Funding:", "Teaching:", "Vision Mātauranga:", "Iwi/hapū & stakeholder engagement and public participation:", 
        "Awards & Honours:", "Opportunities & Issues:", "Other:"]

    # create a new master document
    master, masterFN = create_new_master_docx(year, qtr)
    # get all the docx files in the current directory
    docxFiles = get_docx_files(fileLocation)

    # Add title to the master
    p = master.add_heading(f'JCDR Quaterly Report {year} {qtr}', 0)
    p.alignment = 1

    # loop through all the sections
    for section in range(len(sections)):
        # Add section title to the master
        p = master.add_heading(sections[section], 1)

        # loop through all the docx files
        for doc in docxFiles:
            # load the document
            document = load_docx(fileLocation, doc)
            # get the person name from the document
            # person = document.tables[0].cell(0, 0).text
            person = doc.split('.')[0]
            # append the person name and the first cell of the table to the master
            try:
                append_pragraph_to_docx(master, person, document, section)
            except Exception:
                print(f'Error in {doc}')
                # check if the error folder exists
                if not os.path.exists(f'{fileLocation}Error'):
                    # create the error folder
                    os.mkdir(f'{fileLocation}Error')
                # move this file to the error folder
                os.rename(fileLocation + doc, f'{fileLocation}Error/{doc}')
                # remove this file from the docxFiles list
                docxFiles.remove(doc)
        # Add a seperator
        master.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
        master.add_paragraph()
    # save the master
    save_master_docx(fileLocation, masterFN, master)